import re
from docx import Document

# 打开现有的未规范的 docx 文件
input_file = "学生学习相关的100条QA.docx"
doc = Document(input_file)

# 创建一个新的Word文档对象
output_doc = Document()

# 从docx文件中读取内容
def read_content(doc):
    content = ""
    for para in doc.paragraphs:
        content += para.text + "\n"
    return content

# Markdown 和 HTML 转 Word 格式化函数
def markdown_html_to_word(content, doc):
    # 处理标题 (支持 Markdown 和 HTML 风格)
    content = re.sub(r"^(#+)\s*(.*)", lambda m: f'<h{len(m.group(1))}>{m.group(2)}</h{len(m.group(1))}>', content, flags=re.MULTILINE)
    content = re.sub(r"<h\d>(.*?)<\/h\d>", lambda m: f'<h1>{m.group(1)}</h1>', content)

    # 处理加粗和斜体 (支持 Markdown 和 HTML 风格)
    content = re.sub(r"\*\*(.*?)\*\*|__(.*?)__", r"<b>\1\2</b>", content)
    content = re.sub(r"<b>(.*?)<\/b>", r"<b>\1</b>", content)
    content = re.sub(r"\*(.*?)\*|_(.*?)_", r"<i>\1\2</i>", content)
    content = re.sub(r"<i>(.*?)<\/i>", r"<i>\1</i>", content)

    # 处理无序列表
    ul_list = re.compile(r"^(\*|\+|-) (.+)$", re.MULTILINE)
    content = ul_list.sub(r"<li>\2</li>", content)

    # 将Markdown和HTML文本转换成Word文档格式
    lines = content.split('\n')
    for line in lines:
        if '<h' in line:
            level = int(line[2])
            text = re.sub(r"<\/?h\d>", "", line)
            doc.add_heading(text, level=level)
        elif '<b>' in line:
            para = doc.add_paragraph()
            text = re.sub(r"<\/?b>", "", line)
            run = para.add_run(text)
            run.bold = True
        elif '<i>' in line:
            para = doc.add_paragraph()
            text = re.sub(r"<\/?i>", "", line)
            run = para.add_run(text)
            run.italic = True
        elif '<li>' in line:
            text = re.sub(r"<\/?li>", "", line)
            doc.add_paragraph(text, style='ListBullet')
        else:
            doc.add_paragraph(line)

# 读取内容
content = read_content(doc)

# 转换Markdown和HTML格式到Word
markdown_html_to_word(content, output_doc)

# 保存格式化后的文档
output_file = "Markdown和HTML格式转Word文档.docx"
output_doc.save(output_file)

print(f"文档已保存为 {output_file}")
